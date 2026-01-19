"""
Enterprise RAG System - DOCX Document Processor
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Any, BinaryIO, Optional, Union

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from src.core.exceptions import DocumentProcessingError
from src.core.types import ContentType
from src.ingestion.processors.base import DocumentProcessor, ProcessedDocument


class DocxProcessor(DocumentProcessor):
    """Processor for Microsoft Word documents."""

    supported_extensions = [".docx"]
    supported_mimetypes = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]

    async def process(
        self,
        source: Union[str, Path, BinaryIO],
        metadata: Optional[dict[str, Any]] = None,
    ) -> ProcessedDocument:
        """Process a DOCX file and extract text content."""
        try:
            if isinstance(source, (str, Path)):
                doc = DocxDocument(source)
            else:
                content = source.read()
                doc = DocxDocument(io.BytesIO(content))
        except PackageNotFoundError as e:
            raise DocumentProcessingError(
                document_id="unknown",
                reason=f"Invalid DOCX file: {e}",
            )

        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading structure
                if para.style and para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "")
                    try:
                        level_num = int(level)
                        prefix = "#" * level_num + " "
                    except ValueError:
                        prefix = "## "
                    paragraphs.append(f"{prefix}{text}")
                else:
                    paragraphs.append(text)

        # Extract text from tables
        table_texts = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(cells))
            if table_rows:
                table_texts.append("\n".join(table_rows))

        # Combine all content
        content_parts = paragraphs
        if table_texts:
            content_parts.append("\n--- Tables ---\n")
            content_parts.extend(table_texts)

        content = "\n\n".join(content_parts)

        # Extract document properties
        doc_metadata = {}
        core_props = doc.core_properties
        if core_props.title:
            doc_metadata["docx_title"] = core_props.title
        if core_props.author:
            doc_metadata["docx_author"] = core_props.author
        if core_props.subject:
            doc_metadata["docx_subject"] = core_props.subject
        if core_props.created:
            doc_metadata["docx_created"] = str(core_props.created)
        if core_props.modified:
            doc_metadata["docx_modified"] = str(core_props.modified)

        # Merge with provided metadata
        combined_metadata = {**doc_metadata, **(metadata or {})}
        combined_metadata["paragraph_count"] = len(paragraphs)
        combined_metadata["table_count"] = len(doc.tables)

        # Use document title or first heading
        title = doc_metadata.get("docx_title")
        if not title:
            for para in doc.paragraphs[:5]:
                if para.style and para.style.name.startswith("Heading"):
                    title = para.text.strip()[:100]
                    break
        if not title and paragraphs:
            title = paragraphs[0][:100].lstrip("#").strip()

        return ProcessedDocument(
            content=content,
            content_type=ContentType.TEXT,
            metadata=combined_metadata,
            title=title,
        )
